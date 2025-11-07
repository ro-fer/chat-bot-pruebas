from flask import Flask, request, jsonify, render_template, send_from_directory, Response
import os
from docx import Document
import requests
import logging
import re

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
DOCUMENTS_DIR = "documents"

# ================================
# CONFIGURACIÓN BÁSICA
# ================================
def check_auth(username, password):
    return username == os.environ.get('DOC_USER', 'admin') and password == os.environ.get('DOC_PASS', 'password123')

def authenticate():
    return Response('Acceso requerido', 401, {'WWW-Authenticate': 'Basic realm="Documentación Privada"'})

@app.route('/documentos/<path:filename>')
def download_document(filename):
    auth = request.authorization
    if not auth or not check_auth(auth.username, auth.password):
        return authenticate()
    return send_from_directory(DOCUMENTS_DIR, filename)

@app.route('/documentos/')
def list_documents():
    auth = request.authorization
    if not auth or not check_auth(auth.username, auth.password):
        return authenticate()
    
    documentos = [archivo for archivo in os.listdir(DOCUMENTS_DIR) if archivo.lower().endswith('.docx')]
    html = "<h1>📁 Documentos Disponibles</h1><ul>"
    for doc in documentos:
        html += f'<li><a href="/documentos/{doc}" download>{doc}</a></li>'
    html += "</ul><p><em>Usa Ctrl+Click para descargar</em></p>"
    return html

# ================================
# DETECTOR DE CONTEXTO DE CHAT
# ================================
def es_chat_flotante(request):
    """Detecta si la solicitud viene del widget flotante"""
    # Método 1: Por referer o headers
    referer = request.headers.get('Referer', '')
    if 'probando-widget' in referer:
        return True
    
    # Método 2: Por parámetro en la solicitud
    if request.json and request.json.get('source') == 'widget':
        return True
        
    return False

# ================================
# PROCESADOR DE DOCX MEJORADO
# ================================
def procesar_docx_completo(ruta_archivo):
    """Procesa TODO el contenido del DOCX de forma más completa"""
    try:
        doc = Document(ruta_archivo)
        texto_completo = ""
        
        # Procesar párrafos con mejor estructura
        for paragraph in doc.paragraphs:
            texto = paragraph.text.strip()
            if texto:
                # Identificar títulos importantes
                if any(keyword in texto.lower() for keyword in 
                      ['equipo de', 'dirección', 'coordinación', 'analistas', 'objetivos', 'actividades']):
                    texto_completo += f"\n\n{texto}\n"
                else:
                    texto_completo += f"{texto}\n"
        
        # Procesar tablas de forma más inteligente
        for table in doc.tables:
            texto_completo += "\n" + "="*50 + "\n"
            for i, row in enumerate(table.rows):
                fila_texto = []
                for cell in row.cells:
                    if cell.text.strip():
                        fila_texto.append(cell.text.strip())
                if fila_texto:
                    # Mejor formato para tablas
                    if i == 0:  # Encabezado de tabla
                        texto_completo += " | ".join(fila_texto) + "\n"
                        texto_completo += "-" * 30 + "\n"
                    else:
                        texto_completo += " | ".join(fila_texto) + "\n"
            texto_completo += "="*50 + "\n\n"
        
        return texto_completo.strip()
        
    except Exception as e:
        logger.error(f"Error procesando DOCX {ruta_archivo}: {str(e)}")
        return f"ERROR: {str(e)}"

def extraer_secciones_especificas(contenido):
    """Extrae y enfatiza secciones específicas que sabemos que existen"""
    secciones_importantes = []
    
    # Buscar específicamente el Equipo de Imagen
    lineas = contenido.split('\n')
    en_equipo_imagen = False
    equipo_imagen_content = []
    
    for i, linea in enumerate(lineas):
        linea_limpia = linea.strip()
        
        # Detectar inicio del Equipo de Imagen
        if 'equipo de imagen' in linea_limpia.lower():
            en_equipo_imagen = True
            equipo_imagen_content.append(f"\n🎨 **EQUIPO DE IMAGEN - INFORMACIÓN COMPLETA**\n")
            continue
        
        # Si estamos en la sección de imagen, capturar contenido
        if en_equipo_imagen:
            # Detectar fin de sección (nuevo equipo)
            if i > 0 and any(otro_equipo in linea_limpia.lower() for otro_equipo in 
                            ['equipo de', 'equipo de monitoreo', 'equipo de proyectos']):
                if 'imagen' not in linea_limpia.lower():
                    break
            
            # Agregar contenido relevante
            if len(linea_limpia) > 5:
                equipo_imagen_content.append(linea_limpia)
    
    # Si encontramos información del equipo de imagen, agregarla al inicio
    if equipo_imagen_content:
        secciones_importantes.append("\n".join(equipo_imagen_content))
    
    return secciones_importantes

def cargar_documentos_docx():
    documentos = {}
    if not os.path.exists(DOCUMENTS_DIR):
        logger.warning(f"Directorio {DOCUMENTS_DIR} no existe")
        return documentos
    
    archivos = os.listdir(DOCUMENTS_DIR)
    logger.info(f"Archivos en directorio: {archivos}")
    
    for archivo in archivos:
        if archivo.lower().endswith('.docx'):
            ruta_archivo = os.path.join(DOCUMENTS_DIR, archivo)
            texto_base = procesar_docx_completo(ruta_archivo)
            
            if texto_base and not texto_base.startswith("ERROR"):
                # Mejorar el contenido con secciones específicas
                secciones_especiales = extraer_secciones_especificas(texto_base)
                
                # Combinar contenido base con secciones especiales
                contenido_mejorado = texto_base
                if secciones_especiales:
                    contenido_mejorado = "\n\n".join(secciones_especiales) + "\n\n" + contenido_mejorado
                    logger.info(f"✅ Secciones especiales agregadas a {archivo}")
                
                documentos[archivo] = contenido_mejorado
                logger.info(f"✅ Documento {archivo} cargado y mejorado")
            else:
                logger.error(f"❌ Error cargando {archivo}: {texto_base}")
    
    return documentos

# ================================
# BÚSQUEDA LOCAL MEJORADA
# ================================
def buscar_localmente(pregunta, documentos):
    """Búsqueda local mejorada para cuando Groq no está disponible"""
    pregunta_limpia = pregunta.lower()
    
    # Diccionario de equipos y términos
    equipos = {
        'stock': ['stock', 'equipamiento', 'inventario'],
        'proyectos': ['proyectos', 'implementación', 'analistas de proyectos'],
        'soporte': ['soporte', 'técnico', 'tic', 'instalación'],
        'imagen': ['imagen', 'cartelería', 'señalética'],
        'monitoreo': ['monitoreo', 'vinculación'],
        'dirección': ['dirección', 'programa']
    }
    
    # Términos específicos
    terminos_especificos = {
        'reequipamiento': ['reequipamiento', 'recambio de equipamiento'],
        'instalación': ['instalación', 'instalaciones técnicas'],
        'cartelería': ['cartelería', 'señalética'],
        'inauguración': ['inauguración', 'ceremonia']
    }
    
    resultados = []
    
    for doc_nombre, contenido in documentos.items():
        # Buscar por equipos
        for equipo, palabras in equipos.items():
            if any(palabra in pregunta_limpia for palabra in palabras):
                # Buscar sección del equipo
                lineas = contenido.split('\n')
                en_seccion = False
                
                for i, linea in enumerate(lineas):
                    linea_limpia = linea.strip()
                    
                    if any(palabra in linea_limpia.lower() for palabra in palabras):
                        if not en_seccion:
                            resultados.append(f"<strong>🏢 {equipo.upper()}</strong><br>")
                            en_seccion = True
                        
                        # Capturar contexto
                        inicio = max(0, i-1)
                        fin = min(len(lineas), i+6)
                        for j in range(inicio, fin):
                            if lineas[j].strip() and len(lineas[j].strip()) > 5:
                                resultados.append(f"• {lineas[j].strip()}<br>")
                        break
                break
        
        # Buscar por términos específicos
        if not resultados:
            for termino, palabras in terminos_especificos.items():
                if any(palabra in pregunta_limpia for palabra in palabras):
                    lineas = contenido.split('\n')
                    for i, linea in enumerate(lineas):
                        if any(palabra in linea.lower() for palabra in palabras) and len(linea.strip()) > 10:
                            resultados.append(f"<strong>🔍 INFORMACIÓN SOBRE {termino.upper()}</strong><br>")
                            inicio = max(0, i-1)
                            fin = min(len(lineas), i+4)
                            for j in range(inicio, fin):
                                if lineas[j].strip():
                                    resultados.append(f"• {lineas[j].strip()}<br>")
                            break
                    break
        
        # Búsqueda general si no encontró nada específico
        if not resultados:
            lineas = contenido.split('\n')
            for i, linea in enumerate(lineas):
                if pregunta_limpia in linea.lower() and len(linea.strip()) > 10:
                    resultados.append(f"<strong>📄 {doc_nombre}</strong><br>")
                    resultados.append(f"<strong>🔍 INFORMACIÓN RELACIONADA:</strong><br>")
                    inicio = max(0, i-1)
                    fin = min(len(lineas), i+4)
                    for j in range(inicio, fin):
                        if lineas[j].strip():
                            resultados.append(f"• {lineas[j].strip()}<br>")
                    break
    
    if resultados:
        return "".join(resultados)
    
    return f"""
    🤔 <strong>No encontré información específica sobre "{pregunta}"</strong><br><br>
    
    💡 <strong>Prueba con:</strong><br>
    • <strong>"Stock"</strong> - Equipamiento e inventario<br>
    • <strong>"Proyectos"</strong> - Implementación y gestión<br>
    • <strong>"Soporte técnico"</strong> - Instalación y mantenimiento<br>
    • <strong>"Instalación"</strong> - Procesos técnicos<br>
    • <strong>"Reequipamiento"</strong> - Cambio de equipamiento<br>
    • <strong>"Cartelería"</strong> - Imagen y señalética<br>
    • <strong>"Puesta en marcha"</strong> - Procedimientos completos<br>
    """

# ================================
# GROQ - MEJOR CONTEXTO Y MÁS ESTRICTO
# ================================
def preguntar_groq(pregunta, documentos):
    api_key = os.environ.get('GROQ_API_KEY')
    
    if not api_key:
        return buscar_localmente(pregunta, documentos)

    try:
        # Construir contexto más estricto
        contexto = "INFORMACIÓN EXACTA DEL DOCUMENTO - USAR SOLO ESTO:\n\n"
        
        for doc_nombre, contenido in documentos.items():
            contexto += f"=== DOCUMENTO: {doc_nombre} ===\n"
            
            # Tomar contenido completo pero limitar tamaño
            lineas = contenido.split('\n')
            lineas_importantes = []
            
            for linea in lineas:
                # Filtrar líneas con contenido sustancial
                if (len(linea.strip()) > 10 and 
                    not linea.startswith('===') and
                    not linea.startswith('---')):
                    lineas_importantes.append(linea.strip())
                    if len(lineas_importantes) >= 60:  # Más líneas para mejor contexto
                        break
            
            contexto += "\n".join(lineas_importantes) + "\n\n"
            
            if len(contexto) > 10000:
                contexto += "[... contenido adicional disponible ...]\n\n"
                break
        
        # System prompt MÁS ESTRICTO
        system_prompt = """Eres un asistente que SOLO puede usar la información proporcionada. 
REGLAS ESTRICTAS:
1. NO inventes información
2. NO supongas nada  
3. NO agregues conocimiento externo
4. Si no hay información suficiente, di que no la tienes
5. Usa SOLO el texto proporcionado

Responde en español con HTML básico: <br> para saltos, <strong>para negritas</strong>."""

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": "llama-3.1-8b-instant",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"INFORMACIÓN DISPONIBLE (usar SOLO esto):\n{contexto}\n\nPREGUNTA: {pregunta}\n\nRESPUESTA (HTML básico, usar SOLO información proporcionada):"}
            ],
            "temperature": 0.1,  # Muy bajo para evitar invención
            "max_tokens": 800,
            "top_p": 0.3  # Más restrictivo
        }
        
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=25
        )
        
        if response.status_code == 200:
            data = response.json()
            respuesta = data["choices"][0]["message"]["content"]
            
            # Mejorar formato HTML
            if '<br>' not in respuesta:
                respuesta = respuesta.replace('\n', '<br>')
            
            return respuesta
            
        elif response.status_code == 429:
            return buscar_localmente(pregunta, documentos)
            
        else:
            return buscar_localmente(pregunta, documentos)
            
    except Exception as e:
        logger.error(f"Error Groq: {str(e)}")
        return buscar_localmente(pregunta, documentos)

# ================================
# RUTAS PRINCIPALES
# ================================
@app.route('/')
def home():
    return render_template('chat.html')

@app.route('/probando-widget')
def nueva_pagina():
    """Tu nueva página con el widget flotante"""
    return render_template('index.html')

@app.route('/api/chat', methods=['POST'])
def chat():
    try:
        data = request.json
        pregunta = data.get('prompt', '').strip()
        es_widget = es_chat_flotante(request)  # ← DETECTAMOS EL CONTEXTO
        
        if not pregunta:
            return jsonify({'success': False, 'error': 'Escribe una pregunta'})
        
        documentos = cargar_documentos_docx()
        
        if not documentos:
            return jsonify({
                'success': True, 
                'response': "📂 No hay documentos en la carpeta 'documents'."
            })
        
        # RESPUESTAS RÁPIDAS DIFERENCIADAS
        pregunta_lower = pregunta.lower()
        
        # Saludo inicial - DIFERENTE SEGÚN EL CONTEXTO
        if any(s in pregunta_lower for s in ['hola', 'buenos días', 'buenas', '/start']):
            if es_widget:
                # Respuesta para WIDGET (con sugerencias de botones)
                return jsonify({
                    'success': True, 
                    'response': f"""¡Hola! 👋 <strong>Tina - Asistente Puntos Digitales</strong><br><br>
💡 <strong>Puedes preguntar sobre:</strong><br>
• <strong>Stock</strong> - Equipamiento e inventario<br>
• <strong>Proyectos</strong> - Implementación y gestión<br>  
• <strong>Soporte técnico</strong> - Instalación y mantenimiento<br>
• <strong>Instalación</strong> - Procesos técnicos<br>
• <strong>Cartelería</strong> - Imagen y señalética<br><br>
📚 <em>Documentos cargados: {len(documentos)}</em>""",
                    'widget_mode': True  # ← Nuevo flag para el frontend
                })
            else:
                # Respuesta para CHAT NORMAL
                return jsonify({
                    'success': True, 
                    'response': f"¡Hola! 👋 Asistente especializado en Puntos Digitales<br><br>📚 Documentos cargados: {len(documentos)}<br>¿En qué puedo ayudarte?"
                })
        
        if any(s in pregunta_lower for s in ['chao', 'adiós', 'bye']):
            return jsonify({'success': True, 'response': "¡Hasta luego! 👋"})
        
        if any(p in pregunta_lower for p in ['documento', 'archivo', 'disponible']):
            docs = list(documentos.keys())
            doc_list = "<br>".join([f"• {d}" for d in docs])
            return jsonify({
                'success': True,
                'response': f"<strong>📂 Documentos ({len(docs)}):</strong><br>{doc_list}"
            })
        
        # Procesar pregunta normal con Groq
        respuesta = preguntar_groq(pregunta, documentos)
        
        # Si es widget, agregar formato especial
        if es_widget:
            return jsonify({
                'success': True, 
                'response': respuesta,
                'widget_mode': True  # ← Indicar al frontend que es widget
            })
        else:
            return jsonify({
                'success': True, 
                'response': respuesta
            })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Error: {str(e)}'})

# ================================
# INICIO
# ================================
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    documentos = cargar_documentos_docx()
    print(f"🚀 ChatBot Puntos Digitales - {len(documentos)} documentos")
    print("✅ Sistema dual: Chat normal + Widget flotante")
    print("📍 Chat normal: http://localhost:5000")
    print("📍 Widget flotante: http://localhost:5000/probando-widget")
    app.run(host='0.0.0.0', port=port, debug=False)